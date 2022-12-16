import pandas as pd
import os
from pathlib import Path
import openpyxl
from openpyxl.styles import PatternFill
import shutil
from docx import Document
import datetime
import inspect

"""
Compares an old and new operational plan files and generates a copy of an operational plan file that highlights the changed cells as blue and shows the old and new values.

Methods
----------
run(old_operational_plan_path, new_operational_plan_path) -> None
compare_opertional_plan_files(old_operational_plan_path, new_operational_plan_path) -> None
are_2_excel_same(excel_1_path, excel_2_path) -> bool
open_file(file_path) -> None
generate_report(changes: dict) -> None:
"""

def run(old_operational_plan_path, new_operational_plan_path) -> None:
    print("Comparing: -------------------------------------")
    print(old_operational_plan_path)
    print(new_operational_plan_path)

    compare_opertional_plan_files(old_operational_plan_path, new_operational_plan_path)

def compare_opertional_plan_files(old_operational_plan_path, new_operational_plan_path) -> None:

    are_they_same = are_2_excel_same(old_operational_plan_path, new_operational_plan_path)
    print("Are they same?", are_they_same)
    if are_they_same:
        print("Comparison skipped")
        return
    else:
        print("Comparison started")

    result_file_path = get_result_file_path()
    shutil.copyfile(old_operational_plan_path, result_file_path)
    changes_dict = generate_change_report_xlsx(old_operational_plan_path, new_operational_plan_path, result_file_path)
    open_file(result_file_path)
    generate_change_report_docx(changes_dict)

def get_result_file_path():
    result_folder_path = Path(__file__).parent.parent / "output" # /AIRE/output
    os.makedirs(result_folder_path, exist_ok=True)
    result_file_path = result_folder_path / "Operational_Plan_Changes.xlsx" # /AIRE/output/Operational_Plan_Changes.xlsx
    return result_file_path

def are_2_excel_same(excel_1_path, excel_2_path) -> bool:
    are_same = False

    df1 = pd.read_excel(excel_1_path)
    df2 = pd.read_excel(excel_2_path)

    if df1.equals(df2):
        are_same = True
    
    return are_same

def open_file(file_path) -> None:
    os.system("start " + str(file_path))

def generate_change_report_xlsx(old_operational_plan_path, new_operational_plan_path, result_file_path):
    sheet_names = ["RUN", "GROW", "TRANSFORM", "RUN-CLUSTER SERVICES", "INTAKE", "STRATEGY IMPLEMENTATION - SDWG", "Definitions"]
    changes_dict = {"RUN":[], "GROW":[], "TRANSFORM":[], "RUN-CLUSTER SERVICES":[], "INTAKE":[], "STRATEGY IMPLEMENTATION - SDWG":[], "Definitions":[]}

    old_workbook = openpyxl.load_workbook(old_operational_plan_path)
    new_workbook = openpyxl.load_workbook(new_operational_plan_path)
    result_workbook = openpyxl.load_workbook(result_file_path)
    for sheet_name in sheet_names:
        old_sheet = old_workbook[sheet_name]
        new_sheet = new_workbook[sheet_name]
        result_sheet = result_workbook[sheet_name]
        for row in range(1, result_sheet.max_row):
            for col in range(1, result_sheet.max_column):
                if old_sheet.cell(row, col).value != new_sheet.cell(row, col).value:
                    changes_dict[sheet_name].append((row, new_sheet.cell(row, 1).value, new_sheet.cell(1, col).value, old_sheet.cell(row, col).value, new_sheet.cell(row, col).value))
                    result_sheet.cell(row, col).value = '{} --> {}'.format(old_sheet.cell(row, col).value, new_sheet.cell(row, col).value)
                    result_sheet.cell(row, col).fill = PatternFill(start_color="A8F3FF", fill_type = "solid")
    result_workbook.save(result_file_path)

    return changes_dict

def generate_change_report_docx(changes: dict) -> None:

    now = datetime.datetime.now().strftime("%m/%d/%Y")

    document = Document()
    document.add_heading("Operational Plan Changes Report")
    document.add_paragraph(now)

    change_num = 0
    for key, values in changes.items():
        if len(values) > 0:
            for value in values:
                change_num += 1
    str = "Change number: {0}".format(change_num)
    document.add_heading(str, level=3)

    for key, values in changes.items():
        if len(values) > 0:
            str = "[{0}]".format(key)
            document.add_heading(str, level=3)
            for value in values:
                str = "Row Number: {0}\tProject: {1}\nColumn: {2}\nValue change: {3} -> {4}".format(value[0], value[1], value[2], value[3], value[4])
                document.add_paragraph(str)

    file_path = Path(__file__).parent.parent / "output" /"Operational_Plan_Changes.docx" # /AIRE/output/Operational_Plan_Changes.xlsx
    document.save(file_path)  