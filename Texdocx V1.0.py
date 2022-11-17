from docx import Document
from docx.shared import Pt
import PySimpleGUI as sg
import re
import os
# Developer: Muhammad Hamza Sohail, email: hamza.sohail29@gmail.com
# Before using this program 
# first create a folder on desktop
# then move all the files in it 
# make sure the scrip_file.py is saved in the folder
# open vs code and open the script file through that folder
# then run the script 
# gui opens select folder and run it 
# type the footer if not needed leave blank
# exit the program


def Text_Replacer(document, word , replace):

    for p in document.paragraphs:
        if word.search(p.text):
            inline = p.runs

            for i in range(len(inline)):
                if word.search(inline[i].text):
                    text = word.sub(replace, inline[i].text)
                    inline[i].text = text
    doc = document             
    style_f = doc.styles['Footer']
    font_f = style_f.font
    font_f.name = 'Arial'
    font_f.bold = False
    font_f.size = Pt(10)
    footer = doc.sections[0].footer
    paragraph_f = footer.paragraphs[0]
    paragraph_f.text = "FRO-018 (NOV 08, 2022)                                              © Kings’s Printer for Ontario, 2022" # insert new value here.
    paragraph_f.style = doc.styles['Footer']# this is what changes the style



def Text_Replacer_Para(document, word , replace):

    for p in document.paragraphs:
        if word.search(p.text):
            inline = p.runs

            for i in range(len(inline)):
                if word.search(inline[i].text):
                    text = word.sub(replace, inline[i].text)
                    inline[i].text = text






def batch_processor1(array,replace, word):
    for i in array:
        if i != "text_script.py":
            words = re.compile(replace)
            replace = word
            filename = i
            doc = Document(filename)
            Text_Replacer(doc, words, replace)
            doc.save(filename)
# version 1.0


def Text_Replacer_Footer(document,string,string2,string3):
    doc = document             
    style_f = doc.styles['Footer']
    font_f = style_f.font
    font_f.name = 'Arial'
    font_f.bold = False
    font_f.size = Pt(10)
    footer = doc.sections[0].footer
    paragraph_f = footer.paragraphs[0]
    
    # every "\t" from left footer to right footer
    s=""
    s2=""
    s3=""
    s=string
    s2=string2
    s3=string3
    word=s+"\t"+s2+"\t"+s3
    paragraph_f.text = word
    paragraph_f.style = doc.styles['Footer']# this is what changes the style




# folder path
def path_adder(dir_path):
    # list to store files
    res = []

    # Iterate directory
    for path in os.listdir(dir_path):
        # check if current path is a file
        if os.path.isfile(os.path.join(dir_path, path)):
            res.append(path)
    return res

    #paragraph_f.text = "FRO-018 (NOV 08, 2022)\t© Kings’s Printer for Ontario, 2022" # insert new value here.


#applying the footer in all the files
def batch_processor(array,string,string2,string3,path):
    for i in array:
        if i != "text_script.py":
            filename = path+"\\"+i
            doc = Document(filename)
            Text_Replacer_Footer(doc,string,string2,string3)
            doc.save(filename)


# GUI Part
layout = [
    [sg.Text("Choose the Folder on the desktop: "),sg.Input(key="-IN-"),sg.FolderBrowse()],
    [sg.Text("Type the left footer of the file:     "),sg.Input(key="-left-")],
    [sg.Text("Type the middle footer of the file:"),sg.Input(key="-mid-")],
    [sg.Text("Type the right footer of the file:   "),sg.Input(key="-right-")],
    [sg.Exit(), sg.Button("Convert the Files"),sg.Text("\t\t\t\t© EAO for Ontario, 2022")],

]
window=sg.Window("\t\t\t\tText Automation Tool", layout, [30,30])
while True:
    event, values = window.read()
    print(event,values)
    if event in (sg.WINDOW_CLOSED,"Exit"):
       break
    if event =="Convert the Files":
        path_directory = values["-IN-"]
        path_directory1 =path_directory
        left = values["-left-"]
        mid = values["-mid-"]
        right= values["-right-"]
        res = path_adder(path_directory)
        
        batch_processor(res,left,mid,right,path_directory)
window.close()