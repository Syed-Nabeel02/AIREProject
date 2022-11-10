
from docx import Document
from docx.shared import Pt
import re
import os


def Text_Replacer(document, word , replace):

    for p in document.paragraphs:
        if word.search(p.text):
            inline = p.runs

            for i in range(len(inline)):
                if word.search(inline[i].text):
                    text = word.sub(replace, inline[i].text)
                    inline[i].text = text
    doc = document             
    style_f = doc.styles['Footer']
    font_f = style_f.font
    font_f.name = 'Arial'
    font_f.bold = False
    font_f.size = Pt(10)
    footer = doc.sections[0].footer
    paragraph_f = footer.paragraphs[0]
    paragraph_f.text = "FRO-018 (NOV 08, 2022)                                              © Kings’s Printer for Ontario, 2022" # insert new value here.
    paragraph_f.style = doc.styles['Footer']# this is what changes the style



def Text_Replacer_Para(document, word , replace):

    for p in document.paragraphs:
        if word.search(p.text):
            inline = p.runs

            for i in range(len(inline)):
                if word.search(inline[i].text):
                    text = word.sub(replace, inline[i].text)
                    inline[i].text = text


def Text_Replacer_Footer(document, word , replace):
    doc = document             
    style_f = doc.styles['Footer']
    font_f = style_f.font
    font_f.name = 'Arial'
    font_f.bold = False
    font_f.size = Pt(10)
    footer = doc.sections[0].footer
    paragraph_f = footer.paragraphs[0]
    paragraph_f.text = "FRO-018 (NOV 08, 2022)                                              © Kings’s Printer for Ontario, 2022" # insert new value here.
    paragraph_f.style = doc.styles['Footer']# this is what changes the style




# folder path
def path_adder(dir_path):
    # list to store files
    res = []

    # Iterate directory
    for path in os.listdir(dir_path):
        # check if current path is a file
        if os.path.isfile(os.path.join(dir_path, path)):
            res.append(path)
    print(res)
    return res

# directory_path = r'C:\Users\SohailMo\OneDrive - Government of Ontario\Desktop\Text_automate'

def batch_processor1(array,replace, word):
    for i in array:
        if i != "text_script.py":
            words = re.compile(replace)
            replace = word
            filename = i
            doc = Document(filename)
            Text_Replacer(doc, words, replace)
            doc.save(filename)

def batch_processor2(array,replace, word):
    for i in array:
        if i != "text_script.py":
            words = re.compile(replace)
            replace = word
            filename = i
            doc = Document(filename)
            Text_Replacer_Footer(doc, words, replace)
            doc.save(filename)

def batch_processor3(array,replace, word):
    for i in array:
        if i != "text_script.py":
            words = re.compile(replace)
            replace = word
            filename = i
            doc = Document(filename)
            Text_Replacer_Para(doc, words, replace)
            doc.save(filename)


directory_path = r'C:\Users\SohailMo\OneDrive - Government of Ontario\Desktop\\'        
directory_path1= input(r"Please enter your folder name on your desktop: ")
path_directory = (directory_path+directory_path1)
res = path_adder(path_directory)
function=input('What function do you wish to execute: \nBoth (press1),\nfor paragraph (press 2),\nfor  footer (press 3)')
if function == "1":
    replace = input("Please input the word you need to remove:")
    word=input('Please input the word you need to replace it by: ')
    batch_processor1(res,replace,word)
elif function == "2":
    replace = input("Please input the word you need to remove:")
    word=input('Please input the word you need to replace it by: ')
    batch_processor2(res,replace,word)
else:
    replace="Queen"
    word="king"
    batch_processor2(res,replace,word)
