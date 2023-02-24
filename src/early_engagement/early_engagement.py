from pathlib import Path
import shutil
from datetime import datetime
import os
from pprint import pprint 
from docx import Document
from docxtpl import DocxTemplate
import pandas
import openpyxl
from openpyxl.styles import PatternFill
from concurrent.futures import ThreadPoolExecutor

# terms
#
# op: operational plan
# shutil: shell utilities
# os: operating system

class EarlyEngagement:


    path_to_output = Path().absolute() / 'data' / 'output' / 'early_engagement' 
    path_to_intake_form_template = Path().absolute() / 'data' / 'input' / 'early_engagement' / 'EA Engagement Self-Assessment Template v0.6.docx'
    path_to_intake_form_folder = Path().absolute() / 'data' / 'output' / 'early_engagement' / 'intake forms'
    path_to_archive = Path().absolute() / 'data' / 'archive' / 'early_engagement'

    def __init__(self, path_to_current_op):
        """
        1. Receives a path to an operational plan file
        2. Initializes [data] dictionary variable
        3. Put current datetime into [data]
        4. Copy the received operational plan file to the output folder
        5. Put the path to the operational plan file in the output folder into [data]
        """
        print("__init__ starts")

        if path_to_current_op == None:
            raise Exception('!!! __init__: No path_to_current_op argument given')

        self.data = {}

        self.data['current_datetime'] = self.get_current_datetime()
        self.data['original_path_to_current_op'] = path_to_current_op

        shutil.copyfile(path_to_current_op, self.path_to_output / "current_op.xlsx")
        self.data['path_to_current_op'] = self.path_to_output / "current_op.xlsx"

        print("__init__ finished")

    def check_first_run(self):
        """
        1. Check whether there was a previous run by checking the archive folder
        """
        print("check_first_run starts")

        if(self.is_first_run()):
            self.data['is_first_run'] = True
        else:
            self.data['is_first_run'] = False

        print("check_first_run finished")

        return self;

    def add_previous_op_to_output_folder(self):
        """
        1. Copy the previous operational plan file in the archive folder into the output folder
        2. Add the path to the previous operational plan file in the output folder to [data]
        """
        print("add_previous_op_to_output_folder starts")

        if(self.data['is_first_run'] == True):
            return self

        shutil.copyfile(self.path_to_archive / "previous_op.xlsx", self.path_to_output / "previous_op.xlsx")
        self.data['path_to_previous_op'] = self.path_to_output / "previous_op.xlsx"

        print("add_previous_op_to_output_folder finished")

        return self

   
    def compare_current_previous_op(self):
        """
        1. Compare the previous and current operational plans 
        2. Initializes [comparison] dictionary variable
        2. Put comparison results (same or not, changed cell location and values if there is any) into [comparison]
        3. Put [comparison] into [data]
        """
        print("compare_current_previous_op starts")

        if(self.data['is_first_run']):
            return self

        comparison = {}

        comparison['are_run_same'] = self.are_previous_current_sheets_same('RUN')
        comparison['are_grow_same'] = self.are_previous_current_sheets_same('GROW')
        comparison['are_transform_same'] = self.are_previous_current_sheets_same('TRANSFORM')

        if(comparison['are_run_same'] == False):
            comparison['changes_in_run_sheet'] = self.compare_previous_current_sheets_of('RUN')
        if(comparison['are_grow_same'] == False):
            comparison['changes_in_grow_sheet'] = self.compare_previous_current_sheets_of('GROW')
        if(comparison['are_transform_same'] == False):
            comparison['changes_in_transform_sheet'] = self.compare_previous_current_sheets_of('TRANSFORM')

        self.data['comparison'] = comparison

        print("compare_current_previous_op finished")

        return self
        
    def generate_comparison_report(self):
        """
        1. Generate a Word Document object
        2. Iterate through [data] to put information into the Word Document object
        3. Save the Word Document object as a Word file in the output folder.
        """
        print("generate_comparison_report starts")

        if(self.data['is_first_run']):
            return self

        doc = Document()
        for key, value in self.data.items():
            doc.add_paragraph(f'{key}: {value}')
            doc.add_paragraph("--------------------------------------------------------------------------------")
        comparison_report_name = "Comparison_Report_" + self.data['current_datetime'].replace('-', '').replace(':','').replace(' ','_') + ".docx"
        doc.save(self.path_to_output / comparison_report_name)

        print("generate_comparison_report finished")

        return self

    def generate_comparison_tables(self):
        """
        1. Make a copy of the current operational plan file to use it to describe the changed cells
        2. If one of the sheets (RUN/GROW/TRANSFORM) are changed, compare the previous and current operational plan files of the changed sheet, highlight the changed cells,and write both previous and current values in the copied file
        """
        print("generate_comparison_table starts")

        if(self.data['is_first_run']):
            return self

        comparison_tables_name = "Comparison_Tables_" + self.data['current_datetime'].replace('-', '').replace(':','').replace(' ','_') + ".xlsx"
        path_to_comparison_tables = self.path_to_output / comparison_tables_name
        shutil.copyfile(self.data['path_to_current_op'], path_to_comparison_tables)

        if(self.data['comparison']['are_run_same'] == False):
            self.generate_comparison_table_of('RUN', path_to_comparison_tables)
        if(self.data['comparison']['are_grow_same'] == False):
            self.generate_comparison_table_of('GROW', path_to_comparison_tables)
        if(self.data['comparison']['are_transform_same'] == False):
            self.generate_comparison_table_of('TRANSFORM', path_to_comparison_tables)

        print("generate_comparison_table finished")

        return self

    def generate_intake_forms(self):
        print("generate_intake_forms starts")

        if not os.path.exists(self.path_to_intake_form_folder):
            os.makedirs(self.path_to_intake_form_folder)

        for sheet_name in ["RUN", "GROW", "TRANSFORM"]:
            dataframe = pandas.read_excel(self.data['path_to_current_op'], sheet_name)
            intake_form_count = len(dataframe.index)
            for index, record in enumerate(dataframe.to_dict(orient="records")[:]): # for each operational plan item
                self.generate_intake_form(record, sheet_name, index, intake_form_count)

        print("generate_intake_forms finished")

        return self;

    def archive_files(self):
        """
        1. Make a zip file of the output folder and put it in the data/archive folder
        2. Save the current operational plan file in the data/archive folder for the next run (will be used as a previous operational plan)
        """
        print("archive_files starts")
        
        shutil.copyfile(self.path_to_output / "current_op.xlsx", self.path_to_archive / "previous_op.xlsx")
   
        archive_name = self.format_datetime_string(self.data['current_datetime'])
        shutil.make_archive(self.path_to_archive / archive_name, 'zip', self.path_to_output)

        print("archive_files finished")
     
        return self

    def clear_output_folder(self):
        """
        1. Delete all the files in the output folder for the next run
        """
        print("clear_output_folder starts")

        shutil.rmtree(self.path_to_output)
        os.makedirs(self.path_to_output)

        print("clear_output_folder finished")

        return self

    

    # ----------------------------- helper methods -----------------------------

    def is_first_run(self):
        if not os.path.exists(self.path_to_archive):
            return True
        if len(os.listdir(self.path_to_archive)) == 0:
            return True
        return False

    def get_current_datetime(self):
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def format_datetime_string(self, datetime_string):
        # replace - and : to solve file name format error
        return datetime_string.replace('-', '').replace(':','').replace(' ','_') 

    def are_previous_current_sheets_same(self, sheet_name):
        previous_sheet_dataframe = pandas.read_excel(self.data['path_to_previous_op'], sheet_name=sheet_name)
        current_sheet_dataframe = pandas.read_excel(self.data['path_to_current_op'], sheet_name=sheet_name)
        return previous_sheet_dataframe.equals(current_sheet_dataframe)

    def compare_previous_current_sheets_of(self, sheet_name):
        changes = []

        previous_op_workbook = openpyxl.load_workbook(self.data['path_to_previous_op'])
        current_op_workbook = openpyxl.load_workbook(self.data['path_to_current_op'])

        previous_sheet = previous_op_workbook[sheet_name]
        current_sheet = current_op_workbook[sheet_name]

        for row in range(1, current_sheet.max_row):
            for col in range(1, current_sheet.max_column):
                if previous_sheet.cell(row, col).value != current_sheet.cell(row, col).value:
                    changes.append((row, col, previous_sheet.cell(row, col).value, current_sheet.cell(row, col).value))

        return changes

    def generate_comparison_table_of(self, sheet_name, path_to_comparison_tables):

        previous_op_workbook = openpyxl.load_workbook(self.data['path_to_previous_op'])
        current_op_workbook = openpyxl.load_workbook(self.data['path_to_current_op'])
        comparison_result_workbook = openpyxl.load_workbook(path_to_comparison_tables)

        previous_sheet = previous_op_workbook[sheet_name]
        current_sheet = current_op_workbook[sheet_name]
        comparison_result_sheet = comparison_result_workbook[sheet_name]

        for row in range(1, current_sheet.max_row):
            for col in range(1, current_sheet.max_column):
                if previous_sheet.cell(row, col).value != current_sheet.cell(row, col).value:
                    comparison_result_sheet.cell(row, col).value = '{} --> {}'.format(previous_sheet.cell(row, col).value, current_sheet.cell(row, col).value)
                    comparison_result_sheet.cell(row, col).fill = PatternFill(start_color="A8F3FF", fill_type = "solid")

        comparison_result_workbook.save(path_to_comparison_tables)

        return self

    def generate_intake_form(self, record, sheet_name, index, intake_form_count):
        print(int(index)+1, "/", intake_form_count,  "started")

        doc = DocxTemplate(self.path_to_intake_form_template)
        intake_form_name = self.generate_intake_form_name(record, sheet_name)
        doc = self.fill_in_intake_form(doc, record, intake_form_name)

        path_to_intake_form = self.path_to_intake_form_folder / intake_form_name

        if not os.path.isfile(path_to_intake_form):
            intake_form_name = "NEW_" + intake_form_name
            path_to_intake_form = self.path_to_intake_form_folder / intake_form_name
            doc.save(path_to_intake_form) 

    def generate_intake_form_name(self, record: dict, sheet_name: str) -> str:
        """
        1. Generate an intake form file name based on the item's details and sheet_name
        """
        IDColumn = record['ID']
        InitColumn = record['Initiative']
        ItemCol = record['WorkItemName']
        BranchColumn = record['AccountableBranch'][-6:]

        if(isinstance(IDColumn, str)):
            InitColumn = InitColumn.strip()
        if(isinstance(ItemCol, str)):
            ItemCol = ItemCol.strip()
        if(isinstance(BranchColumn, str)):
            BranchColumn = BranchColumn.strip()

        if record['MustDoCantFail'] == 'Yes':
            intake_form_name = f"{IDColumn}{'_'}{BranchColumn}{'_'}{'MDCF'}{'_'}{sheet_name[0]}{'_'}{InitColumn}{'_'}{ItemCol}"
        else:
            intake_form_name = f"{IDColumn}{'_'}{BranchColumn}{'_'}{sheet_name[0]}{'_'}{InitColumn}{'_'}{ItemCol}"
        
        intake_form_name = intake_form_name.replace(' ', '').replace('-', '').replace('.', '') + ".docx"
        
        return intake_form_name

    def fill_in_intake_form(self, doc: DocxTemplate, record:dict, intake_form_name: str) -> DocxTemplate:
        """
        1. Fill in the template with received item's details
        2. Record current datetime and intake form name for tracking (solution branches might change intake forms' name)
        """
        doc.render(record) 
        doc.add_paragraph("This was Autogenerated on" + " " + self.data['current_datetime'])
        doc.add_paragraph("Do not modify:" + intake_form_name)
        return doc

if __name__ == '__main__':
    # For testing purpose
    path_to_current_op = Path().absolute() / 'data' / 'input' / 'early_engagement' / 'CYSSC FY 2022-23 Operational Plan - PUBLISHED June 2022.xlsx'
    path_to_changed_op = Path().absolute() / 'data' / 'input' / 'early_engagement' / 'CYSSC FY 2022-23 Operational Plan - PUBLISHED June 2022 - old.xlsx'
    
    try:
        EarlyEngagement(path_to_current_op).check_first_run().add_previous_op_to_output_folder().compare_current_previous_op().generate_comparison_report().generate_comparison_tables().generate_intake_forms().archive_files().clear_output_folder()
    except Exception as e:
        print(e)