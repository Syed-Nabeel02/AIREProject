import openai
import requests
import base64
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
import os
import glob

openai.api_key = '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'

# Path to your image directory and image name
image_directory = "C:\\AIRE\\data\\input\\image_description" #Change path
image_name = "ImageTest3" #Change name

# Define the output directory
output_dir = "C:\\AIRE\\data\\output\\image_description" #path

def main():
    # Generate a report based on the contents of the provided ZIP file
    encode_image(image_name, image_directory)

# Function to encode the image
def encode_image(image_name, directory):
    # Search for the file in the specified directory
    for file in glob.glob(f"{directory}\\{image_name}.*"):
        # Check if the file is an image
        if file.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif')):
            with open(file, "rb") as image_file:
                return base64.b64encode(image_file.read()).decode('utf-8'), file
        else:
            raise ValueError("The input is not an Image")
    return None, None

# Getting the base64 string and image path
base64_image, image_path = encode_image(image_name, image_directory)

headers = {
    "Content-Type": "application/json",
    "Authorization": f"Bearer {openai.api_key}"
}

payload = {
    "model": "gpt-4-vision-preview",
    "messages": [
      {
        "role": "user",
        "content": [
          {
            "type": "text",
            "text": "What’s in this image?"
          },
          {
            "type": "image_url",
            "image_url": {
              "url": f"data:image/jpeg;base64,{base64_image}"
            }
          }
        ]
      }
    ],
    "max_tokens": 600
}

response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)

# Create a new Word document
doc = Document()

# Add the 'Image:' text in red
image_paragraph = doc.add_paragraph()
image_run = image_paragraph.add_run('Image:')
font = image_run.font
font.color.rgb = RGBColor(255, 0, 0)  # RGB values for red

# Add the image with increased size
if image_path is not None:
    doc.add_picture(image_path, width=Inches(6)) # Adjust the size as needed
else:
    print("No image file found.")

# Add a line of space
doc.add_paragraph()

# Add the 'Image Translation:' text in red
translation_paragraph = doc.add_paragraph()
translation_run = translation_paragraph.add_run('Image Translation:')
font = translation_run.font
font.color.rgb = RGBColor(255, 0, 0)  # RGB values for red

if 'choices' in response.json() and response.json()['choices']:
    text_description = response.json()['choices'][0]['message']['content']
    # Add the description to the document
    doc.add_paragraph(text_description)
else:
    print("No text description found in the response.")

# Check if the directory exists, if not, create it
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Save the document in the output directory
doc.save(os.path.join(output_dir, 'Image_Description.docx'))

if __name__ == "__main__":
    main()
