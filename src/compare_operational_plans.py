import pandas as pd
from docx import Document
from pathlib import Path
import os

new_file = Path().absolute() / 'data' / 'input' / 'compare_operational_plans' / 'new_operational_plan.xlsx'
old_file = Path().absolute() / 'data' / 'input' / 'compare_operational_plans' / 'new_operational_plan_old.xlsx'
output_file = Path().absolute() / 'data' / 'output' / 'compare_operational_plans' / 'Operational Plan Comparison Report.docx'
path_to_output = Path().absolute() / 'data' / 'output' / 'compare_operational_plans'

def main():
    print(f"Comparing {new_file} with {old_file}...")
    # Start the comparison process
    compare_excels(new_file, old_file, output_file)
    open_output_folder();

def read_excel(file_path):
    # Reads an Excel file and returns a DataFrame
    return pd.read_excel(file_path)

def get_changes(row_new, row_old):
    # Compare two rows and return changes as a dictionary
    changes = {}
    for column, value in row_new.items():
        old_value = row_old.get(column)
        if value != old_value and pd.notna(value) and pd.notna(old_value):
            changes[column] = (old_value, value)
    return changes

# Following functions are responsible for writing different sections of the comparison report
def write_modified_items(doc, modified_items):
    # Writes modified items to the Word document
    doc.add_heading('1. Modified Items', level=2)
    doc.add_paragraph(f'Number of Modified Items: {len(modified_items)}')
    for index, item in enumerate(modified_items, start=1):
        para = doc.add_paragraph()
        run = para.add_run(f"{index}. Modified Item: ({item['item_id']}) {item['item_name']}")
        run.bold = True
        for change_index, (column, change) in enumerate(item['changes'].items(), start=1):
            old_value, new_value = change
            para.add_run(f'\n    {change_index}. Change in [{column}]: [{old_value}] -> [{new_value}]')

def write_new_items(doc, new_items):
    # Writes new items to the Word document
    doc.add_heading('2. New Items', level=2)
    doc.add_paragraph(f'Number of New Items: {len(new_items)}')
    for index, (item_id, item_name) in enumerate(new_items, start=1):
        doc.add_paragraph(f"{index}. New Item: ({item_id}) {item_name}")

def write_deleted_items(doc, deleted_items):
    # Writes deleted items to the Word document
    doc.add_heading('3. Deleted Items', level=2)
    doc.add_paragraph(f'Number of Deleted Items: {len(deleted_items)}')
    for index, (item_id, item_name) in enumerate(deleted_items, start=1):
        doc.add_paragraph(f"{index}. Deleted Item: ({item_id}) {item_name}")

def log_counts(modified_items, new_items, deleted_items):
    # Prints the counts of modified, new, and deleted items to the console
    print(f"Number of Modified Items: {len(modified_items)}")
    print(f"Number of New Items: {len(new_items)}")
    print(f"Number of Deleted Items: {len(deleted_items)}")

def compare_excels(new_file, old_file, output_file):
    # Main function to compare two Excel files and write differences to a Word document
    # Read Excel files
    df_new = read_excel(new_file)
    df_old = read_excel(old_file)

    # Create the directory if it does not exist
    output_file.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading('Operational Plan Comparison Report', level=1)

    modified_items, new_items, deleted_items = process_items(df_new, df_old)

    log_counts(modified_items, new_items, deleted_items)

    write_modified_items(doc, modified_items)
    write_new_items(doc, new_items)
    write_deleted_items(doc, deleted_items)

    doc.save(output_file)
    print(f"Differences written to {output_file}.")

def process_items(df_new, df_old):
    # Compares two DataFrames and returns lists of modified, new, and deleted items
    modified_items = []
    new_items = []
    deleted_items = []

    # Loop through new items and compare with old items
    for index, row_new in df_new.iterrows():
        item_name = row_new['Item Name']
        accountable_branch = row_new['Accountable Branch']
        item_id = row_new.get('Item ID') if pd.notna(row_new.get('Item ID')) else ''
        matching_rows_old = df_old[(df_old['Item Name'] == item_name) & (df_old['Accountable Branch'] == accountable_branch)]

        if not matching_rows_old.empty:
            row_old = matching_rows_old.iloc[0]
            changes = get_changes(row_new, row_old)
            if changes:
                modified_item = {"item_name": item_name, "item_id": item_id, "changes": changes}
                modified_items.append(modified_item)
        else:
            new_items.append((item_id, item_name))

    # Find deleted items by looping through old items and comparing with new items
    for index, row_old in df_old.iterrows():
        item_name = row_old['Item Name']
        accountable_branch = row_old['Accountable Branch']
        item_id = row_old.get('Item ID') if pd.notna(row_old.get('Item ID')) else ''
        if not ((df_new['Item Name'] == item_name) & (df_new['Accountable Branch'] == accountable_branch)).any():
            deleted_items.append((item_id, item_name))

    return modified_items, new_items, deleted_items

def open_output_folder() -> None:
    os.startfile(str(path_to_output))

if __name__ == "__main__":
    main()
